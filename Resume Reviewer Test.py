#from resume_parser import resumeparse
from pyresparser import ResumeParser
from docx import Document

import nltk
nltk.download('punkt')
from nltk import sent_tokenize, word_tokenize

filed = input()
doc = Document()

try:
    with open(filed, 'r') as file:
        doc.add_paragraph(file.read())
    doc.save("text.docx")
    data = ResumeParser('text.docx').get_extracted_data()
    print(data['name'])
except:
    data = ResumeParser(filed).get_extracted_data()
    print(data['name'])

#print(data)

'''#from resume_parser import resumeparse
from pyresparser import ResumeParser
from docx import Document

#file = ('Srushti Vaidyanathan - College Resume.docx')

##file format should be in .txt , .docx or .pdf only
filed = input()

try:
    doc = Document()
    with open(filed, 'r') as file:
        doc.add_paragraph(file.read())
    doc.save("text.docx")
    data = ResumeParser('text.docx').get_extracted_data()
    print(data['name'])
except:
    data = ResumeParser(filed).get_extracted_data()
    print(data['name'])

print(data['name'])

#name = data['name'].word'''



