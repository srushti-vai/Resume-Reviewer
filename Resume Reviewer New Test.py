from docx import Document

import nltk
nltk.download('punkt')
from nltk import sent_tokenize, word_tokenize

nltk.download('stopwords')
from nltk.corpus import stopwords

nltk.download('wordnet')
from nltk.stem import PorterStemmer, WordNetLemmatizer
ps = PorterStemmer()
wn = WordNetLemmatizer()
from nltk.corpus import wordnet

nltk.download('averaged_perceptron_tagger')
nltk.download('tagsets')

from nltk.chunk import RegexpParser

# ------------------------------------------------------------------------- #

#document = Document('C:\Srushti\College Admissions\Purdue\LaunchPad\Resume-Reviewer\Test Resumes\Srushti Vaidyanathan - College Resume.docx')
document = Document(input('Input your resume document in .docx format!\n'))

# ------------------------------------------------------------------------- #

# gets the set of repeating verbs in resume
def repeating_words():
    words = []

    for para in document.paragraphs:
        line = nltk.word_tokenize(para.text)
        for l in line:
            words.append(l)


    #stores the words from the doc that aren't stop words
    stop_words = set(stopwords.words('english'))

    wo_stop_word_tokens = []
    for word in words:
        if word not in stop_words:
            wo_stop_word_tokens.append(word)

    # ------------------------------------------------------------------------- #
    #stores the words without stop words lemmatized
    """ lemmatized_words = []

    for word in wo_stop_word_tokens:
        lemmatized_words.append(wn.lemmatize(word))
    """

    # ------------------------------------------------------------------------- #
    #stores the lemmatized words without stop words with their part of speech (pos)
    pos_words = nltk.pos_tag(wo_stop_word_tokens)

    #print(pos_words)

    # ------------------------------------------------------------------------- #
    # gets only the verbs from all of the words in the document
    
    temp_verbs = []

    for word, pos in pos_words:
        if pos.find('VB') > -1: 
            temp_verbs.append(word)
    
    # combs through the words twice because sometimes it doesn't catch everything the first time

    temp_verbs = nltk.pos_tag(temp_verbs)
    verbs = []

    for word, pos in temp_verbs:
        if pos.find('VB') > -1:  # 'NN' represents a noun
            verbs.append(word)

    # -------------------------------------------------------------------------- #
    # finds if words are duplicates
    
    duplicate_verbs = []

    for i in range(len(verbs)):
        if verbs[i] in verbs[i+1:-1] and verbs[i] not in duplicate_verbs:
            duplicate_verbs.append(verbs[i])

    # -------------------------------------------------------------------------- #
    # suggests synonyms for the duplicate verbs

    synonyms = {}

    for verb in duplicate_verbs:
        synonyms[verb] = []
        for syn in wordnet.synsets(verb):
            for lemma in syn.lemmas():
                synonyms[verb].append(lemma.name())

    # ------------------------------------------------------------------------- #

    print('The following verbs are repeated in your resume a few times:')
    print(*synonyms.keys())
    cont = input('Would you like to replace any of these duplicate verbs?\n');
    if (cont.lower() == 'yes'):
       while True:
            which_verb = input("Which verb would you like to fix?\n")
            if which_verb in synonyms.keys():
                print('These are possible synonyms:')
                print(synonyms[which_verb])
                select = input('Select a word that fits!\n')
                synonyms[which_verb].remove(select)
                print('Good choice!')
                print(synonyms[which_verb])
            else:
                print('Please spell the word correctly.')
                continue
            proceed = input('Would you like to continue?\n')
            if (proceed.lower() == 'no'):
                break


        

repeating_words()
